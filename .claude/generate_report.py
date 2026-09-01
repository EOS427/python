# -*- coding: utf-8 -*-
"""生成《跨编辑器 Python 运行环境差异》报告 docx，保存到桌面。"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DESKTOP = r"d:\桌面"
OUT = DESKTOP + r"\跨编辑器Python运行环境差异报告.docx"

doc = Document()

# 设置默认字体为中文字体
def set_font(run, size=10.5, bold=False, mono=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    if mono:
        run.font.name = "Consolas"
    else:
        run.font.name = "Calibri"
    # 中文字体
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), "微软雅黑" if not mono else "Consolas")
    if color:
        run.font.color.rgb = RGBColor(*color)

def title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_font(r, size=20, bold=True, color=(0x1F, 0x3B, 0x63))
    return p

def h1(text):
    p = doc.add_heading(level=1)
    r = p.add_run(text)
    set_font(r, size=16, bold=True, color=(0x1F, 0x3B, 0x63))
    return p

def h2(text):
    p = doc.add_heading(level=2)
    r = p.add_run(text)
    set_font(r, size=13, bold=True, color=(0x2E, 0x54, 0x8A))
    return p

def h3(text):
    p = doc.add_heading(level=3)
    r = p.add_run(text)
    set_font(r, size=11.5, bold=True, color=(0x40, 0x40, 0x40))
    return p

def body(text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, size=10.5, bold=bold)
    p.paragraph_format.space_after = Pt(4)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r0 = p.add_run(bold_prefix)
        set_font(r0, size=10.5, bold=True)
    r = p.add_run(text)
    set_font(r, size=10.5)
    p.paragraph_format.space_after = Pt(2)
    return p

def code(lines):
    for line in lines.split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.25)
        # 浅灰底纹
        pPr = p._element.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:fill'), 'F2F2F2')
        pPr.append(shd)
        r = p.add_run(line if line else " ")
        set_font(r, size=9.5, mono=True, color=(0x20, 0x20, 0x20))

# ======================= 正文 =======================
title("跨编辑器 Python 运行环境差异导致的导入错误\n底层原理与解决方案报告")

body("项目：python testing prograhm　|　涉及文件：pycharm/深度学习实践/复现mnist/rebuild.py　|　日期：2026-09-01", bold=True)

# ---------- 一、问题背景 ----------
h1("一、问题背景")
body("同一份代码 rebuild.py，在 PyCharm 中运行时，from common.functions import *、from dataset.mnist import load_mnist 以及 open(...sample_weight.pkl) 都能正常工作；"
     "但换到 Trae（一款基于 VS Code 内核的 AI 编辑器）中，同样的代码却报出 ModuleNotFoundError（找不到 common / dataset 模块）或 FileNotFoundError（找不到 sample_weight.pkl 文件）。")
body("代码本身没有语法错误，真正的原因在于：不同编辑器为 Python 解释器「补充模块搜索路径」的机制不同，以及代码里使用了依赖「当前工作目录」的相对路径。"
     "本报告用通俗语言逐层拆解背后的底层逻辑，并给出可落地的解决方案。")

# ---------- 二、核心概念 ----------
h1("二、先搞懂三个底层概念")
body("要理解这个 Bug，只需要三个概念：模块搜索路径（sys.path）、当前工作目录（CWD）、以及编辑器如何“偷偷”修改这两者。")

h2("2.1 Python 找模块的依据：sys.path")
body("当 Python 执行 import 语句时，它并不在整个硬盘上搜索，而是拿着一个「目录清单」，按顺序在里面找。这个清单就是 sys.path，它是一个字符串列表，"
     "每个元素是一个目录的绝对路径。")
bullet("import common.functions 这句话的真实含义是：", "核心结论：")
code("在 sys.path 里的每一个目录下，找有没有一个叫 common 的包（目录），里面有没有 functions.py。\n"
     "任何一个目录里找到了，就成功；全部找完都没有，就抛 ModuleNotFoundError。")
body("因此，想让 from common.functions import * 成功，唯一条件就是：包含 common 的那一层目录，必须出现在 sys.path 里。")

h3("sys.path 的初始化顺序（以“直接运行脚本”为例）")
bullet("脚本文件所在的目录（运行时，它会被放在 sys.path 的第 0 位）。")
bullet("环境变量 PYTHONPATH 里写的所有目录。")
bullet("Python 标准库目录（如 Lib、site-packages 的上级）。")
bullet("第三方包目录 site-packages。")
body("代码里写的 sys.path.append(...) 属于「运行时手动追加」，会加到这个清单的末尾（或插入指定位置）。")

h2("2.2 相对路径与当前工作目录（CWD）")
body("open(\"a/b.pkl\")、open(\"../x/y.pkl\") 这类「相对路径」，是相对于「当前工作目录」（Current Working Directory，简称 CWD）来解析的，"
     "而不是相对于脚本文件所在的目录。")
bullet("CWD 是什么：", "要点：")
body("CWD 是操作系统给每个运行中的进程记录的一个“当前所在文件夹”。你在哪个目录下启动了这个进程，CWD 通常就是哪个目录。")
bullet("在 PyCharm / Trae / 命令行里运行同一个脚本，CWD 可能各不相同，于是同一个相对路径会指向不同位置——这正是“在这台编辑器能跑，换个编辑器就不行”的常见根源之一。", "")
bullet("os.pardir 只是字符串 \"..\"（表示上一级目录），sys.path.append(os.pardir) 加进去的其实是「CWD 的上一级」，而且它是相对路径，其实际指向完全取决于 CWD。", "")

h2("2.3 编辑器如何「悄悄」修改 sys.path")
body("这里就是 PyCharm 和 Trae 产生差异的关键。")
h3("PyCharm 的做法：主动“认领”目录")
bullet("Sources Root 机制：", "（1）")
body("PyCharm 允许你把任意目录「标记为 Sources Root（源根目录）」。一旦标记，编辑器在静态分析（也就是代码里有没有红线）时，就把该目录当作 import 的根。")
bullet("运行配置默认「Add source roots to PYTHONPATH」勾选：", "（2）")
body("这是 PyCharm 运行配置里的一个默认开启选项。运行脚本时，PyCharm 会把你在项目里标记的所有 Sources Root 自动塞进 sys.path，再启动解释器。")
body("结论：PyCharm 是“主动型”——它默认帮你把需要的目录注入 sys.path，所以即便代码里的 sys.path.append 没起关键作用，import 依然能成功。")

h3("Trae（VS Code 系）的做法：被动“只信环境”")
body("Trae 使用 Pylance（微软的 Python 语言服务器）做代码分析。Pylance 默认不会主动猜测哪些目录该加入搜索路径，它只相信两样东西：")
bullet("Python 解释器自身的环境（PYTHONPATH 环境变量、site-packages 等）。")
bullet("配置文件 .vscode/settings.json 里显式写明的 python.analysis.extraPaths。")
body("结论：Trae 是“被动型”——你没有显式告诉它 common 在哪，它就报红；运行时也没有额外注入 sys.path，于是真正执行也失败。")

# ---------- 三、具体错误分析 ----------
h1("三、本项目的具体错误拆解")

h2("3.1 根本原因：目录「多嵌套」了一层")
body("GitHub 上下载的 zip 包解压后，往往会得到「外层一个同名文件夹 + 内层一个同名文件夹」的结构。本项目正是这种情况：")
code("python testing prograhm/\n"
     "└── pycharm/\n"
     "    └── 深度学习实践/\n"
     "        ├── 复现mnist/\n"
     "        │   └── rebuild.py                 ← 脚本在这里\n"
     "        └── deep-learning-from-scratch-master/        ← 外层空壳（只有一层内嵌文件夹）\n"
     "            └── deep-learning-from-scratch-master/    ← 真正的仓库根目录\n"
     "                ├── common/functions.py\n"
     "                ├── dataset/mnist.py\n"
     "                └── ch03/sample_weight.pkl")
body("注意：common、dataset、ch03 这三个目标，都在「再往下数两层」的目录里，而不是脚本的上一级就能直接看到。这就是一切问题的根源。")

h2("3.2 错误一：import common / dataset 失败")
code("import sys, os\n"
     "sys.path.append(os.pardir)      # 只把「CWD 的上一级」加进去\n"
     "from common.functions import *   # 报 ModuleNotFoundError\n"
     "from dataset.mnist import load_mnist")
body("假设 CWD 是 复现mnist/，那么 os.pardir 指向 深度学习实践/，但 common、dataset 在 深度学习实践/deep-learning-from-scratch-master/deep-learning-from-scratch-master/ 下面。"
     "距离差了两层，自然找不到。")
body("在 PyCharm 里，因为 Sources Root 机制把真正的那层目录补进了 sys.path，错误被「掩盖」；换到 Trae，没有这个机制，问题就暴露了。")

h2("3.3 错误二：pickle 文件路径错误")
body("原始代码：")
code("with open(\"../deep-learning-from-scratch-master/sample_weight.pkl\", 'rb') as f:")
body("这个路径存在两处偏差：")
bullet("少了中间一层嵌套目录：真实路径里有两次 deep-learning-from-scratch-master，代码里只有一次。", "偏差一：")
bullet("少了 ch03 子目录：sample_weight.pkl 其实在 ch03/ 里，而不是仓库根目录。", "偏差二：")
bullet("它是相对路径，解析结果完全取决于 CWD；CWD 一变，路径指向就变。", "隐患：")

h2("3.4 错误三（潜在）：对 CWD 的隐性依赖")
body("即便把路径写对了，只要它仍是相对路径，且运行时的 CWD 不是脚本所在目录，问题依然会复现。"
     "这正是「换一个编辑器/换一种启动方式就崩」的深层原因。")

# ---------- 四、解决方案 ----------
h1("四、解决方案")

h2("4.1 方案一（推荐）：代码层自定位，一劳永逸")
body("核心思想：不再依赖 CWD，也不再依赖任何编辑器的隐式配置，而是用脚本自身的绝对位置来推算所有需要路径。")
body("Python 里有个内置变量 __file__，它保存当前脚本文件的路径（可能是相对路径）。配合 os.path.abspath 可以把它变成绝对路径，"
     "再用 os.path.dirname 取所在目录，就得到了脚本的“家”。")
code("import sys, os\n"
     "base = os.path.dirname(os.path.abspath(__file__))   # 脚本所在目录 = .../复现mnist/\n"
     "repo = os.path.join(base, \"..\", \"deep-learning-from-scratch-master\",\n"
     "                    \"deep-learning-from-scratch-master\")   # 真正的仓库根\n"
     "sys.path.append(repo)                              # 让 common/dataset 可被 import\n"
     "import numpy as np\n"
     "import pickle\n"
     "from common.functions import *\n"
     "from dataset.mnist import load_mnist")
body("读取 pickle 时，同样用 repo 拼接：")
code("def init_network():\n"
     "    with open(os.path.join(repo, \"ch03\", \"sample_weight.pkl\"), 'rb') as f:\n"
     "        network = pickle.load(f)\n"
     "    return network")
body("这样做的优点：无论在 PyCharm、Trae、还是命令行双击/终端运行，行为都完全一致，因为路径永远由脚本位置唯一决定。")

h3("各函数的作用（技术细节）")
bullet("__file__：Python 运行脚本时自动注入的变量，值是脚本文件路径。", "（1）")
bullet("os.path.abspath(path)：把（可能是相对的）路径转成绝对路径，同时做规范化。", "（2）")
bullet("os.path.dirname(path)：取路径里「目录」部分，丢掉最后的文件名。", "（3）")
bullet("os.path.join(a, b, ...)：用当前操作系统的分隔符（Windows 用 \\，Linux 用 /）正确拼接路径，避免手写分隔符出错。", "（4）")

h2("4.2 方案二：只让 Trae 编辑器不报红（静态分析）")
body("在 Trae 打开的项目目录对应的 .vscode/settings.json 里，加 python.analysis.extraPaths，告诉 Pylance 到哪里找模块。")
body("如果 Trae 打开的是项目根目录 python testing prograhm：", bold=True)
code("{\n"
     "    \"python.defaultInterpreterPath\": \"D:\\\\miniconda_x86_64bit\\\\envs\\\\thellmbook\\\\python.exe\",\n"
     "    \"python.analysis.typeCheckingMode\": \"off\",\n"
     "    \"python.analysis.extraPaths\": [\n"
     "        \"${workspaceFolder}/pycharm/深度学习实践/deep-learning-from-scratch-master/deep-learning-from-scratch-master\"\n"
     "    ]\n"
     "}")
body("改完需要「重载窗口」（Ctrl+Shift+P → Python: Clear Cache and Reload Window），否则 Pylance 不会立即刷新。")
body("注意：extraPaths 只解决编辑器里的红波浪线和跳转，不解决运行时真正 import 报错。", bold=True)

h2("4.3 方案三：让 Trae 运行时也能 import")
body("运行（F5）时想让解释器也能找到模块，需要在 .vscode/launch.json 里注入 PYTHONPATH，或指定正确的 cwd。")
code("{\n"
     "    \"version\": \"0.2.0\",\n"
     "    \"configurations\": [\n"
     "        {\n"
     "            \"name\": \"Python: rebuild\",\n"
     "            \"type\": \"python\",\n"
     "            \"request\": \"launch\",\n"
     "            \"program\": \"${workspaceFolder}/pycharm/深度学习实践/复现mnist/rebuild.py\",\n"
     "            \"cwd\": \"${workspaceFolder}/pycharm/深度学习实践/复现mnist\",\n"
     "            \"env\": {\n"
     "                \"PYTHONPATH\": \"${workspaceFolder}/pycharm/深度学习实践/deep-learning-from-scratch-master/deep-learning-from-scratch-master\"\n"
     "            }\n"
     "        }\n"
     "    ]\n"
     "}")

h2("4.4 三种方案对比")
tbl = doc.add_table(rows=1, cols=4)
tbl.style = 'Light Grid Accent 1'
hdr = tbl.rows[0].cells
for i, t in enumerate(["方案", "解决范围", "是否依赖编辑器", "评价"]):
    p = hdr[i].paragraphs[0]
    r = p.add_run(t)
    set_font(r, size=10.5, bold=True)
rows = [
    ("方案一：改代码 __file__ 定位", "编辑器红线 + 运行时 + 任何环境", "否，完全通用", "推荐，一劳永逸"),
    ("方案二：extraPaths", "仅编辑器静态红线", "是，仅 Trae/VS Code", "临时止血，不解决运行"),
    ("方案三：launch.json PYTHONPATH", "仅运行时（需用 F5 调试器）", "是，仅 Trae/VS Code", "配合方案二可临时用"),
]
for row in rows:
    cells = tbl.add_row().cells
    for i, t in enumerate(row):
        p = cells[i].paragraphs[0]
        r = p.add_run(t)
        set_font(r, size=10.5)

# ---------- 五、技术细节深入 ----------
h1("五、技术细节深入")

h2("5.1 为什么相对路径 import 和 open() 的“参考系”不同")
body("import 依赖 sys.path（一组目录清单，通常是绝对路径）；open() 依赖 CWD（操作系统为进程维护的一个“当前目录”）。"
     "两者是两套独立机制，互不替代。所以即便 import 成功，open() 仍然可能失败，反之亦然。")

h2("5.2 静态分析（编辑器红线）与运行时是两回事")
body("Pylance 等语言服务器做的是「静态分析」：它不真正执行代码，而是靠配置（extraPaths）推断模块在哪。"
     "真正执行脚本时，是 Python 解释器按 sys.path 找模块。两者配置通道不同，这也是为什么："
     "编辑器里不报红 ≠ 运行时能跑通。")

h2("5.3 如何自己诊断 sys.path 和 CWD")
body("遇到类似问题，最快的方法是在脚本开头打印两样东西，看真实值：")
code("import sys, os\n"
     "print(\"CWD      =\", os.getcwd())\n"
     "print(\"__file__ =\", __file__)\n"
     "print(\"sys.path =\")\n"
     "for p in sys.path:\n"
     "    print(\"   \", p)")
body("打印出来，一眼就能看出：CWD 在哪、脚本在哪、搜索清单里到底有没有包含 common 的那一层目录。")

h2("5.4 为什么 PyCharm 默认能跑，而“看起来一样”的 Trae 不能")
body("本质是“默认配置的侵入性”不同：")
bullet("PyCharm 默认会「Add source roots to PYTHONPATH」，等于运行时帮你改了 sys.path，属于“越俎代庖”的便利。", "")
bullet("Trae/VS Code 的 Pylance 默认尽量“忠实于”解释器环境，不擅自加目录，属于“所见即所得”的克制。", "")
body("两种理念无所谓对错，但一旦代码依赖了编辑器的隐式行为，迁移到另一个编辑器就会踩坑。所以最佳实践是：代码里显式、自洽地处理路径，不依赖任何编辑器。")

# ---------- 六、最佳实践 ----------
h1("六、通用最佳实践")
bullet("永远不要假设 CWD 是脚本所在目录，需要定位文件时用 os.path.abspath(__file__) 推路径。", "① ")
bullet("不要在代码里写死 ../ 这类相对路径去加载关键资源，改成由 __file__ 计算的绝对路径。", "② ")
bullet("import 的路径整理：把需要 import 的「仓库根目录」显式 append 到 sys.path，而不是只上一级。", "③ ")
bullet("跨编辑器协作时，优先保证代码自洽，其次才是给各编辑器补配置（extraPaths / launch.json）。", "④ ")
bullet("解压 GitHub 项目时，留意是否出现「同名文件夹套同名文件夹」，这常是 import 失败的第一元凶。", "⑤ ")

# ---------- 七、总结 ----------
h1("七、总结")
body("本报告的 Bug 表面是“换个编辑器就报错”，本质是三个底层事实叠加：")
bullet("目录多嵌套了一层（解压 zip 的常见陷阱），导致路径对不上。", "1. ")
bullet("代码用了依赖 CWD 的相对路径（open 和 sys.path.append(os.pardir)）。", "2. ")
bullet("PyCharm 会主动注入 sys.path 掩盖了问题，Trae 不注入于是暴露。", "3. ")
body("解决方案以「代码层用 __file__ 自定位」为最优，因为它让代码在任何编辑器、任何启动方式下都行为一致，"
     "从根上消除了对 CWD 和编辑器隐式配置的依赖。")

doc.save(OUT)
print("已保存到:", OUT)
